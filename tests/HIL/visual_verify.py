"""Office MCP 视觉排版验证脚本

使用本地 Word/Excel COM 自动化导出 PDF，再用 pymupdf 渲染为 PNG，
最后提取文档结构信息与 PNG 并排输出，供视觉比对。

用法: uv run python tests/HIL/visual_verify.py
"""

import os
import sys
import time
import win32com.client
import pymupdf
from pathlib import Path

TEST_DIR = Path(__file__).parent / "test_output"
PNG_DIR = TEST_DIR / "png"
PDF_DIR = TEST_DIR / "pdf"


def ensure_dirs():
    PNG_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)


def docx_to_pdf(docx_path: Path) -> Path:
    """Word COM: .docx → .pdf"""
    pdf_path = PDF_DIR / (docx_path.stem + "_word.pdf")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.ExportAsFixedFormat(
            OutputFileName=str(pdf_path.resolve()),
            ExportFormat=17,  # wdExportFormatPDF
            OpenAfterExport=False,
            OptimizeFor=0,  # wdOptimizeForPrint
            Range=0,  # wdExportAllDocument
        )
        doc.Close(SaveChanges=False)
        print(f"  [OK] Word → PDF: {pdf_path.name}")
        return pdf_path
    finally:
        word.Quit()


def xlsx_to_pdf(xlsx_path: Path) -> Path:
    """Excel COM: .xlsx → .pdf"""
    pdf_path = PDF_DIR / (xlsx_path.stem + "_excel.pdf")
    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    try:
        wb = excel.Workbooks.Open(str(xlsx_path.resolve()))
        for ws in wb.Sheets:
            ws.PageSetup.Orientation = 1  # xlLandscape
            ws.PageSetup.FitToPagesWide = 1
            ws.PageSetup.FitToPagesTall = False
            ws.PageSetup.Zoom = False
        wb.ExportAsFixedFormat(
            Type=0,  # xlTypePDF
            Filename=str(pdf_path.resolve()),
            Quality=0,  # xlQualityStandard
            IncludeDocProperties=True,
            IgnorePrintAreas=False,
            OpenAfterPublish=False,
        )
        wb.Close(SaveChanges=False)
        print(f"  [OK] Excel → PDF: {pdf_path.name}")
        return pdf_path
    finally:
        excel.Quit()


def pdf_to_png(pdf_path: Path) -> list[Path]:
    """pymupdf: PDF 每页 → PNG"""
    doc = pymupdf.open(str(pdf_path))
    png_paths = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=200)
        png_path = PNG_DIR / f"{pdf_path.stem}_page{i+1}.png"
        pix.save(str(png_path))
        png_paths.append(png_path)
        print(f"  [OK] PDF page {i+1} → PNG: {png_path.name} ({pix.width}x{pix.height})")
    doc.close()
    return png_paths


def analyze_docx(docx_path: Path) -> dict:
    """分析 docx 结构（用 python-docx）"""
    from docx import Document
    doc = Document(str(docx_path))

    result = {
        "file": docx_path.name,
        "paragraphs": [],
        "tables": [],
        "images": 0,
    }

    for i, para in enumerate(doc.paragraphs):
        info = {
            "index": i,
            "text": para.text[:50],
            "style": para.style.name,
            "alignment": str(para.alignment) if para.alignment else "None",
        }
        # 检查字体属性
        for run in para.runs:
            if run.bold:
                info["bold"] = True
            if run.italic:
                info["italic"] = True
            if run.font.size:
                info["font_size"] = run.font.size.pt
            if run.font.color and run.font.color.rgb:
                info["font_color"] = str(run.font.color.rgb)
            break  # 只取第一个 run 的属性
        result["paragraphs"].append(info)

    for i, table in enumerate(doc.tables):
        tbl = {"index": i, "rows": len(table.rows), "cols": len(table.columns), "cells": []}
        for row in table.rows:
            for cell in row.cells:
                tbl["cells"].append(cell.text[:20])
        result["tables"].append(tbl)

    # 统计图片
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            result["images"] += 1

    return result


def analyze_xlsx(xlsx_path: Path) -> dict:
    """分析 xlsx 结构（用 openpyxl）"""
    from openpyxl import load_workbook
    wb = load_workbook(str(xlsx_path))

    result = {
        "file": xlsx_path.name,
        "sheets": [],
    }

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_info = {
            "name": sheet_name,
            "dimensions": str(ws.dimensions),
            "merged_cells": [str(m) for m in ws.merged_cells.ranges],
            "data": [],
        }

        for row in ws.iter_rows(max_row=min(ws.max_row, 10), values_only=False):
            row_data = []
            for cell in row:
                cell_info = {"value": cell.value}
                if cell.font and cell.font.bold:
                    cell_info["bold"] = True
                if cell.font and cell.font.color and cell.font.color.rgb:
                    cell_info["color"] = str(cell.font.color.rgb)
                if cell.alignment and cell.alignment.horizontal:
                    cell_info["align"] = cell.alignment.horizontal
                if cell.number_format and cell.number_format != "General":
                    cell_info["fmt"] = cell.number_format
                row_data.append(cell_info)
            sheet_info["data"].append(row_data)

        # 图表
        sheet_info["charts"] = len(ws._charts) if hasattr(ws, "_charts") else 0
        result["sheets"].append(sheet_info)

    return result


def print_analysis(title: str, analysis: dict):
    """打印分析结果"""
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")

    if "paragraphs" in analysis:
        print(f"\n  段落数: {len(analysis['paragraphs'])}")
        print(f"  表格数: {len(analysis['tables'])}")
        print(f"  图片数: {analysis['images']}")
        print()
        for p in analysis["paragraphs"]:
            extras = []
            if p.get("bold"):
                extras.append("加粗")
            if p.get("italic"):
                extras.append("斜体")
            if p.get("font_size"):
                extras.append(f"{p['font_size']}pt")
            if p.get("font_color"):
                extras.append(f"色#{p['font_color']}")
            extras_str = f" [{', '.join(extras)}]" if extras else ""
            print(f"  [{p['index']}] style={p['style']:15s} align={p['alignment']:20s}{extras_str}")
            print(f"       \"{p['text']}\"")

        if analysis["tables"]:
            print()
            for t in analysis["tables"]:
                print(f"  表格[{t['index']}]: {t['rows']}行×{t['cols']}列  样式=Table Grid")

    if "sheets" in analysis:
        for s in analysis["sheets"]:
            print(f"\n  Sheet: {s['name']}  ({s['dimensions']})")
            if s["merged_cells"]:
                print(f"  合并单元格: {', '.join(s['merged_cells'])}")
            print(f"  图表数: {s['charts']}")
            print()
            for row in s["data"]:
                cells = []
                for c in row:
                    v = str(c["value"]) if c["value"] is not None else ""
                    flags = []
                    if c.get("bold"):
                        flags.append("B")
                    if c.get("color"):
                        flags.append(f"#{c['color']}")
                    if c.get("align"):
                        flags.append(c["align"])
                    if c.get("fmt"):
                        flags.append(c["fmt"])
                    flag_str = f"({','.join(flags)})" if flags else ""
                    cells.append(f"{v}{flag_str}")
                print(f"    {'  |  '.join(cells)}")


def generate_report(docx_analysis, xlsx_analysis, png_files):
    """生成验证报告 markdown"""
    report_lines = [
        "# Office MCP 排版视觉验证报告",
        "",
        f"**验证时间**: {time.strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 验证方法",
        "",
        "1. 使用 Office MCP 工具创建带格式的 docx/xlsx 文件",
        "2. 使用本地 Word/Excel COM 自动化导出为 PDF",
        "3. 使用 pymupdf 将 PDF 渲染为 200 DPI PNG 图片",
        "4. 提取文档结构信息与 PNG 并排比对",
        "",
        "## 生成的 PNG 文件",
        "",
    ]

    for p in png_files:
        report_lines.append(f"- `{p.name}`")

    report_lines.extend(["", "## DOCX 结构分析", ""])
    for p in docx_analysis["paragraphs"]:
        extras = []
        if p.get("bold"):
            extras.append("加粗")
        if p.get("italic"):
            extras.append("斜体")
        if p.get("font_size"):
            extras.append(f"{p['font_size']}pt")
        if p.get("font_color"):
            extras.append(f"色#{p['font_color']}")
        extras_str = f" ({', '.join(extras)})" if extras else ""
        report_lines.append(f"- [{p['index']}] **{p['style']}** | align={p['alignment']}{extras_str}")
        report_lines.append(f"  - \"{p['text']}\"")

    if docx_analysis["tables"]:
        report_lines.append("")
        for t in docx_analysis["tables"]:
            report_lines.append(f"- 表格[{t['index']}]: {t['rows']}行×{t['cols']}列")

    report_lines.extend(["", "## XLSX 结构分析", ""])
    for s in xlsx_analysis["sheets"]:
        report_lines.append(f"### Sheet: {s['name']}")
        report_lines.append(f"- 范围: {s['dimensions']}")
        if s["merged_cells"]:
            report_lines.append(f"- 合并: {', '.join(s['merged_cells'])}")
        report_lines.append(f"- 图表: {s['charts']}")
        report_lines.append("")

        for row in s["data"]:
            cells = []
            for c in row:
                v = str(c["value"]) if c["value"] is not None else ""
                flags = []
                if c.get("bold"):
                    flags.append("B")
                if c.get("color"):
                    flags.append(f"#{c['color']}")
                if c.get("align"):
                    flags.append(c["align"])
                flag_str = f" ({','.join(flags)})" if flags else ""
                cells.append(f"{v}{flag_str}")
            report_lines.append(f"| {'  |  '.join(cells)} |")

    report_lines.extend([
        "",
        "## 排版验证结论",
        "",
        "| 验证项 | 预期 | 实际 | 状态 |",
        "|--------|------|------|------|",
        "| Title 样式 | 加粗26pt居中 | 待PNG确认 | ⏳ |",
        "| Heading 1 | 16pt | 待PNG确认 | ⏳ |",
        "| Heading 2 | 14pt | 待PNG确认 | ⏳ |",
        "| Heading 3 | 12pt | 待PNG确认 | ⏳ |",
        "| 右对齐段落 | alignment=RIGHT | 待PNG确认 | ⏳ |",
        "| 表格Table Grid | 3×4带边框 | 待PNG确认 | ⏳ |",
        "| Excel表头 | 加粗白色居中 | 待PNG确认 | ⏳ |",
        "| Excel合计行 | 加粗红色 | 待PNG确认 | ⏳ |",
        "| Excel数值格式 | 0.0 | 待PNG确认 | ⏳ |",
        "| Excel图表 | 柱状图 | 待PNG确认 | ⏳ |",
        "",
        "> 请查看 `tests/HIL/test_output/png/` 目录下的 PNG 文件进行视觉比对。",
    ])

    report_path = Path(__file__).parent / "VISUAL_VERIFY_报告.md"
    report_path.write_text("\n".join(report_lines), encoding="utf-8")
    print(f"\n[OK] 报告已生成: {report_path.name}")
    return report_path


def main():
    ensure_dirs()

    docx_file = TEST_DIR / "format_test.docx"
    xlsx_file = TEST_DIR / "format_test.xlsx"

    if not docx_file.exists():
        print(f"[ERROR] 找不到 {docx_file}")
        sys.exit(1)
    if not xlsx_file.exists():
        print(f"[ERROR] 找不到 {xlsx_file}")
        sys.exit(1)

    all_pngs = []

    # --- DOCX ---
    print(f"\n[1/4] 分析 DOCX 结构...")
    docx_analysis = analyze_docx(docx_file)
    print_analysis("DOCX 结构分析", docx_analysis)

    print(f"\n[2/4] Word COM 导出 PDF...")
    docx_pdf = docx_to_pdf(docx_file)
    print(f"\n[3/4] pymupdf 渲染 PNG...")
    docx_pngs = pdf_to_png(docx_pdf)
    all_pngs.extend(docx_pngs)

    # --- XLSX ---
    print(f"\n[2/4] 分析 XLSX 结构...")
    xlsx_analysis = analyze_xlsx(xlsx_file)
    print_analysis("XLSX 结构分析", xlsx_analysis)

    print(f"\n[3/4] Excel COM 导出 PDF...")
    xlsx_pdf = xlsx_to_pdf(xlsx_file)
    print(f"\n[4/4] pymupdf 渲染 PNG...")
    xlsx_pngs = pdf_to_png(xlsx_pdf)
    all_pngs.extend(xlsx_pngs)

    # --- 报告 ---
    print(f"\n[生成报告]")
    generate_report(docx_analysis, xlsx_analysis, all_pngs)

    print(f"\n{'='*60}")
    print(f"  全部完成! PNG 文件位于: {PNG_DIR}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
