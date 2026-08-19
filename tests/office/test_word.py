"""Tests for Word tools."""

import base64
import os
import shutil
import tempfile

from carrot_mcp_office.word import (
    inspect as word_inspect,
    insert_para,
    modify_para,
    format_para,
    delete_para,
    insert_table,
    get_table,
    modify_table,
    format_table,
    delete_table,
    insert_image,
    delete_image,
    get_outline,
    get_content,
    _parse_sections,
)


def _parse_content_result(result):
    """Parse list[TextContent|ImageContent] return into (meta_dict, images_list)."""
    import json
    meta = json.loads(result[0].text)
    images = [r for r in result[1:] if r.type == "image"]
    return meta, images


def _cleanup(original_path):
    from carrot_mcp_office.backup import _mirror_path
    mirror = _mirror_path(original_path)
    if mirror.parent.exists():
        shutil.rmtree(mirror.parent, ignore_errors=True)
    d = os.path.dirname(original_path)
    if os.path.exists(d):
        shutil.rmtree(d, ignore_errors=True)


def _docx():
    d = tempfile.mkdtemp(prefix="test_office_")
    return os.path.join(d, "test.docx")


def _create_test_image():
    """Create a minimal 1x1 red PNG file for testing."""
    path = os.path.join(tempfile.gettempdir(), "test_office_img.png")
    # Minimal 1x1 red PNG
    png_data = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
    )
    with open(path, "wb") as f:
        f.write(png_data)
    return path


def test_insert_para_new_file():
    path = _docx()
    try:
        result = insert_para(path, "Hello World")
        assert result["status"] == "ok"
        assert result["text"] == "Hello World"
        assert result["version"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_insert_para_at_index():
    from docx import Document
    path = _docx()
    try:
        insert_para(path, "First")
        insert_para(path, "Third")
        insert_para(path, "Second", index=1)
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["First", "Second", "Third"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_inspect():
    path = _docx()
    try:
        insert_para(path, "Hello")
        insert_table(path, 2, 2)
        result = word_inspect(path)
        assert result["status"] == "ok"
        assert result["total_paragraphs"] >= 1
        assert result["table_count"] == 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_modify_para():
    path = _docx()
    try:
        insert_para(path, "Original")
        result = modify_para(path, 0, "Modified")
        assert result["status"] == "ok"
        assert result["old_text"] == "Original"
        assert result["new_text"] == "Modified"
        assert result["version"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_modify_para_out_of_range():
    path = _docx()
    try:
        insert_para(path, "Hello")
        result = modify_para(path, 5, "Text")
        assert result["status"] == "error"
        assert "out of range" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_format_para():
    path = _docx()
    try:
        insert_para(path, "Hello")
        result = format_para(path, 0, bold=True, italic=True, alignment="center")
        assert result["status"] == "ok"
        assert result["version"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_format_para_invalid_alignment():
    path = _docx()
    try:
        insert_para(path, "Hello")
        result = format_para(path, 0, alignment="invalid")
        assert result["status"] == "error"
        assert "Invalid alignment" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_delete_para():
    path = _docx()
    try:
        insert_para(path, "First")
        insert_para(path, "Second")
        result = delete_para(path, 0)
        assert result["status"] == "ok"
        assert result["version"] >= 1
        info = word_inspect(path)
        assert info["total_paragraphs"] == 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_delete_para_out_of_range():
    path = _docx()
    try:
        insert_para(path, "Hello")
        result = delete_para(path, 5)
        assert result["status"] == "error"
        assert "out of range" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_insert_table():
    path = _docx()
    try:
        result = insert_table(path, 2, 3, [["A", "B", "C"], ["1", "2", "3"]])
        assert result["status"] == "ok"
        assert result["rows"] == 2
        assert result["cols"] == 3
        assert result["version"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_insert_table_dimension_mismatch():
    path = _docx()
    try:
        result = insert_table(path, 2, 2, [["A", "B", "C"]])
        assert result["status"] == "error"
        assert "dimensions" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_modify_table():
    path = _docx()
    try:
        insert_table(path, 2, 2)
        result = modify_table(path, 0, 0, 0, "Hello")
        assert result["status"] == "ok"
        assert result["new_text"] == "Hello"
        assert result["version"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_modify_table_out_of_range():
    path = _docx()
    try:
        insert_table(path, 2, 2)
        result = modify_table(path, 5, 0, 0, "Text")
        assert result["status"] == "error"
        assert "out of range" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_format_table():
    path = _docx()
    try:
        insert_table(path, 2, 2)
        result = format_table(path, 0, style="Table Grid")
        assert result["status"] == "ok"
        assert result["version"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_format_table_invalid_style():
    path = _docx()
    try:
        insert_table(path, 2, 2)
        result = format_table(path, 0, style="Nonexistent Style")
        assert result["status"] == "error"
        assert "not found" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_delete_table():
    path = _docx()
    try:
        insert_table(path, 2, 2)
        result = delete_table(path, 0)
        assert result["status"] == "ok"
        assert result["version"] >= 1
        info = word_inspect(path)
        assert info["table_count"] == 0
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_delete_table_out_of_range():
    path = _docx()
    try:
        insert_para(path, "Placeholder")
        result = delete_table(path, 0)
        assert result["status"] == "error"
        assert "out of range" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_insert_image():
    path = _docx()
    img_path = _create_test_image()
    try:
        result = insert_image(path, img_path)
        assert result["status"] == "ok"
        assert result["version"] >= 1
        info = word_inspect(path)
        assert info["image_count"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)
        if os.path.exists(img_path):
            os.unlink(img_path)


def test_insert_image_at_index():
    path = _docx()
    img_path = _create_test_image()
    try:
        insert_para(path, "First")
        insert_para(path, "Third")
        result = insert_image(path, img_path, index=1)
        assert result["status"] == "ok"
        info = word_inspect(path)
        assert info["total_paragraphs"] == 3
        assert info["image_count"] == 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)
        if os.path.exists(img_path):
            os.unlink(img_path)


def test_insert_image_with_width():
    path = _docx()
    img_path = _create_test_image()
    try:
        result = insert_image(path, img_path, width=2.0)
        assert result["status"] == "ok"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)
        if os.path.exists(img_path):
            os.unlink(img_path)


def test_delete_image():
    path = _docx()
    img_path = _create_test_image()
    try:
        insert_image(path, img_path)
        info = word_inspect(path)
        assert info["image_count"] == 1
        result = delete_image(path, 0)
        assert result["status"] == "ok"
        assert result["version"] >= 1
        info = word_inspect(path)
        assert info["image_count"] == 0
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)
        if os.path.exists(img_path):
            os.unlink(img_path)


def test_delete_image_out_of_range():
    path = _docx()
    try:
        insert_para(path, "Hello")
        result = delete_image(path, 0)
        assert result["status"] == "error"
        assert "out of range" in result["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_insert_table_at_index():
    path = _docx()
    try:
        insert_table(path, 2, 2, [["A", "B"], ["C", "D"]])
        insert_table(path, 1, 1, [["X"]], index=0)
        info = word_inspect(path)
        assert info["table_count"] == 2
        from docx import Document
        doc = Document(path)
        assert doc.tables[0].rows[0].cells[0].text == "X"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_format_para_font_size_color():
    path = _docx()
    try:
        insert_para(path, "Hello")
        result = format_para(path, 0, font_size=14, font_color="FF0000")
        assert result["status"] == "ok"
        assert result["version"] >= 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def _create_heading_docx():
    """Create a test docx with heading hierarchy and content."""
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("Intro text")
    doc.add_heading("Section 1.1", level=2)
    doc.add_paragraph("Section 1.1 content")
    doc.add_heading("Section 1.2", level=2)
    doc.add_paragraph("Section 1.2 content")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("Chapter 2 content")
    doc.add_heading("Section 2.1", level=2)
    doc.add_paragraph("Section 2.1 content")
    doc.save(path)
    return path


def test_get_outline():
    path = _create_heading_docx()
    try:
        result = get_outline(path)
        assert result["status"] == "ok"
        assert result["count"] == 5
        assert len(result["outline"]) == 2
        assert result["outline"][0]["title"] == "Chapter 1"
        assert result["outline"][0]["level"] == 1
        assert len(result["outline"][0]["children"]) == 2
        assert result["outline"][1]["title"] == "Chapter 2"
        flat = result["flat"]
        assert flat[0]["title"] == "Chapter 1"
        assert flat[1]["title"] == "Section 1.1"
        assert flat[1]["parent"] == "Chapter 1"
        assert flat[3]["title"] == "Chapter 2"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_outline_no_headings():
    path = _docx()
    try:
        insert_para(path, "Just text")
        result = get_outline(path)
        assert result["status"] == "ok"
        assert result["count"] == 0
        assert result["outline"] == []
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content():
    path = _create_heading_docx()
    try:
        # Default markdown mode
        result_md = get_content(path, section=[0])
        meta_md, images = _parse_content_result(result_md)
        assert meta_md["status"] == "ok"
        assert meta_md["count"] == 1
        assert "# Chapter 1" in meta_md["markdown"]
        assert "Intro text" in meta_md["markdown"]

        # Structured json mode
        result_json = get_content(path, section=[0], format="json")
        meta_json, _ = _parse_content_result(result_json)
        assert meta_json["status"] == "ok"
        assert meta_json["count"] == 1
        sec = meta_json["sections"][0]
        assert sec["title"] == "Chapter 1"
        assert sec["level"] == 1
        assert sec["paragraph_range"] is not None
        texts = [p["text"] for p in sec["paragraphs"]]
        assert "Intro text" in texts
        assert "Section 1.1 content" in texts
        assert "Section 1.2 content" in texts
        assert "Chapter 2 content" not in texts
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_multiple():
    path = _create_heading_docx()
    try:
        result = get_content(path, section=[1, 3], format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 2
        assert meta["sections"][0]["title"] == "Section 1.1"
        assert meta["sections"][1]["title"] == "Chapter 2"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_out_of_range():
    path = _create_heading_docx()
    try:
        result = get_content(path, section=[99], format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["sections"] == []
        assert "warning" in meta
        assert "99" in meta["warning"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_with_table():
    path = _create_heading_docx()
    try:
        insert_table(path, 2, 2, [["A", "B"], ["C", "D"]])
        result = get_content(path, [0], format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_empty_section():
    """Test getting content for a section that has no content paragraphs between headings."""
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("H1", level=1)
    doc.add_heading("H2", level=1)
    doc.add_heading("H3", level=1)
    doc.save(path)
    try:
        result_md = get_content(path, [1])
        meta_md, _ = _parse_content_result(result_md)
        assert meta_md["status"] == "ok"
        assert "# H2" in meta_md["markdown"]

        result_json = get_content(path, [1], format="json")
        meta_json, _ = _parse_content_result(result_json)
        sec = meta_json["sections"][0]
        assert sec["title"] == "H2"
        assert sec["paragraphs"] == [{"index": 1, "text": "H2", "style": "Heading 1"}]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_int_single():
    """Test get_content with int parameter for single section."""
    path = _create_heading_docx()
    try:
        result = get_content(path, section=0, format="json")
        meta, images = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 1
        sec = meta["sections"][0]
        assert sec["title"] == "Chapter 1"
        assert sec["level"] == 1
        texts = [p["text"] for p in sec["paragraphs"]]
        assert "Intro text" in texts
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_str_single():
    """Test get_content with string parameter for single section."""
    path = _create_heading_docx()
    try:
        result = get_content(path, section="0", format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 1
        assert meta["sections"][0]["title"] == "Chapter 1"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_outline_deep_nesting():
    """Test outline with 3+ levels of nesting."""
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("L1", level=1)
    doc.add_heading("L2", level=2)
    doc.add_heading("L3", level=3)
    doc.add_heading("L3b", level=3)
    doc.add_heading("L2b", level=2)
    doc.add_heading("L1b", level=1)
    doc.save(path)
    try:
        result = get_outline(path)
        assert result["status"] == "ok"
        assert result["count"] == 6
        tree = result["outline"]
        assert len(tree) == 2
        assert tree[0]["title"] == "L1"
        assert len(tree[0]["children"]) == 2
        assert tree[0]["children"][0]["title"] == "L2"
        assert len(tree[0]["children"][0]["children"]) == 2
        assert tree[0]["children"][1]["title"] == "L2b"
        flat = result["flat"]
        assert flat[2]["title"] == "L3"
        assert flat[2]["parent"] == "L2"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_last_section():
    """Test getting content for the last section (no next heading)."""
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("First", level=1)
    doc.add_paragraph("First content")
    doc.add_heading("Last", level=1)
    doc.add_paragraph("Last content A")
    doc.add_paragraph("Last content B")
    doc.save(path)
    try:
        result = get_content(path, [1], format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        sec = meta["sections"][0]
        assert sec["title"] == "Last"
        texts = [p["text"] for p in sec["paragraphs"]]
        assert "Last content A" in texts
        assert "Last content B" in texts
        assert "First content" not in texts
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_outline_empty_heading_text():
    """Test outline with empty heading text."""
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("", level=1)
    doc.add_heading("Has Text", level=2)
    doc.add_heading("", level=1)
    doc.save(path)
    try:
        result = get_outline(path)
        assert result["status"] == "ok"
        assert result["count"] == 3
        assert result["flat"][0]["title"] == ""
        assert result["flat"][1]["title"] == "Has Text"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_with_images():
    """Test content retrieval returns images as ImageContent attachments."""
    img_path = _create_test_image()
    path = _create_heading_docx()
    try:
        insert_image(path, img_path, index=2)
        result = get_content(path, [0], format="json")
        meta, images = _parse_content_result(result)
        assert meta["status"] == "ok"
        sec = meta["sections"][0]
        assert sec["image_count"] >= 1
        assert len(images) >= 1
        assert images[0].type == "image"
        assert images[0].data is not None
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)
        if os.path.exists(img_path):
            os.unlink(img_path)


def test_get_content_range_string():
    """Test range string like '0-2' expands correctly."""
    path = _create_heading_docx()
    try:
        result = get_content(path, ["0-2"], format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 3
        titles = [s["title"] for s in meta["sections"]]
        assert titles == ["Chapter 1", "Section 1.1", "Section 1.2"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_mixed_spec():
    """Test mixed spec like ['0-1', 3] expands correctly."""
    path = _create_heading_docx()
    try:
        result = get_content(path, ["0-1", 3], format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 3
        titles = [s["title"] for s in meta["sections"]]
        assert titles == ["Chapter 1", "Section 1.1", "Chapter 2"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_string_index():
    """Test string index like '3' works the same as int 3."""
    path = _create_heading_docx()
    try:
        result = get_content(path, ["3"], format="json")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 1
        assert meta["sections"][0]["title"] == "Chapter 2"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_markdown_mode():
    """Test markdown mode returns formatted markdown string and no sections dict."""
    path = _create_heading_docx()
    try:
        result = get_content(path, [0], format="markdown")
        meta, images = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 1
        assert "markdown" in meta
        assert "# Chapter 1" in meta["markdown"]
        assert "Intro text" in meta["markdown"]
        assert "sections" not in meta
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_json_mode():
    """Test json mode returns structured sections with indices and styles."""
    path = _create_heading_docx()
    try:
        result = get_content(path, [0], format="json")
        meta, images = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 1
        assert "sections" in meta
        assert "markdown" not in meta
        sec = meta["sections"][0]
        assert sec["title"] == "Chapter 1"
        assert sec["paragraph_range"] == [0, 5]
        assert isinstance(sec["paragraphs"][0], dict)
        assert sec["paragraphs"][0]["text"] == "Chapter 1"
        assert sec["paragraphs"][0]["style"] == "Heading 1"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_paragraph_mode_markdown_and_json():
    """Test get_content with paragraph index in both markdown and json mode."""
    path = _create_heading_docx()
    try:
        # Markdown mode
        res_md = get_content(path, paragraph="1-2")
        meta_md, _ = _parse_content_result(res_md)
        assert meta_md["status"] == "ok"
        assert meta_md["mode"] == "paragraph"
        assert "Intro text" in meta_md["markdown"]

        # JSON mode
        res_json = get_content(path, paragraph="1-2", format="json")
        meta_json, _ = _parse_content_result(res_json)
        assert meta_json["status"] == "ok"
        assert meta_json["mode"] == "paragraph"
        assert isinstance(meta_json["paragraphs"][0], dict)
        assert meta_json["paragraphs"][0]["text"] == "Intro text"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_parse_sections_ints():
    assert _parse_sections([0, 2, 5], 10) == [0, 2, 5]


def test_parse_sections_range():
    assert _parse_sections(["0-3"], 10) == [0, 1, 2, 3]


def test_parse_sections_mixed():
    assert _parse_sections(["0-2", 4, "6-8"], 10) == [0, 1, 2, 4, 6, 7, 8]


def test_parse_sections_dedup():
    assert _parse_sections([0, "0-1"], 10) == [0, 1]


def test_parse_sections_sorted():
    assert _parse_sections([5, "0-2"], 10) == [0, 1, 2, 5]


def test_parse_sections_string_int():
    assert _parse_sections(["3"], 10) == [3]


def test_parse_sections_negative():
    assert _parse_sections(["-1"], 10) == [-1]


def test_parse_sections_int_single():
    assert _parse_sections(0, 5) == [0]


def test_parse_sections_int_middle():
    assert _parse_sections(2, 5) == [2]


def test_parse_sections_int_out_of_range():
    assert _parse_sections(10, 5) == []


def test_parse_sections_str_single():
    assert _parse_sections("0", 5) == [0]


def test_parse_sections_none():
    assert _parse_sections(None, 5) == []


def test_get_table():
    path = os.path.join(tempfile.mkdtemp(), "test_get_table.docx")
    try:
        insert_table(path, rows=3, cols=2, data=[["a", "b"], ["c", "d"], ["e", "f"]])
        # Default markdown format
        result_md = get_table(path, table_index=0)
        assert result_md["status"] == "ok"
        assert result_md["rows"] == 3
        assert result_md["cols"] == 2
        assert "| a | b |" in result_md["markdown"]
        assert "| c | d |" in result_md["markdown"]

        # Explicit json format
        result_json = get_table(path, table_index=0, format="json")
        assert result_json["status"] == "ok"
        assert result_json["data"] == [["a", "b"], ["c", "d"], ["e", "f"]]
    finally:
        _cleanup(path)


def test_get_table_row_range():
    path = os.path.join(tempfile.mkdtemp(), "test_get_table_range.docx")
    try:
        insert_table(
            path,
            rows=4,
            cols=2,
            data=[["Header1", "Header2"], ["row1", "val1"], ["row2", "val2"], ["row3", "val3"]],
        )
        result = get_table(path, table_index=0, row_range="2-3", format="json")
        assert result["status"] == "ok"
        assert result["rows"] == 2
        assert result["data"] == [["row2", "val2"], ["row3", "val3"]]

        # Markdown mode with row_range keeps header at top if not explicitly selected
        result_md = get_table(path, table_index=0, row_range="2-3", format="markdown")
        assert "| Header1 | Header2 |" in result_md["markdown"]
        assert "| row2 | val2 |" in result_md["markdown"]
        assert "row1" not in result_md["markdown"]
    finally:
        _cleanup(path)


def test_get_table_out_of_range():
    path = os.path.join(tempfile.mkdtemp(), "test_get_table_oor.docx")
    try:
        insert_table(path, rows=2, cols=2)
        result = get_table(path, table_index=5)
        assert result["status"] == "error"
        assert "out of range" in result["message"]
    finally:
        _cleanup(path)


# ── Heading lookup and Markdown content tests ──


def test_get_content_by_heading_exact():
    path = _create_heading_docx()
    try:
        # Markdown mode
        result = get_content(path, heading="Chapter 1")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["mode"] == "heading"
        assert meta["count"] == 1
        assert "# Chapter 1" in meta["markdown"]
        assert "Intro text" in meta["markdown"]

        # JSON mode
        result_json = get_content(path, heading="Chapter 1", format="json")
        meta_json, _ = _parse_content_result(result_json)
        assert meta_json["sections"][0]["title"] == "Chapter 1"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_by_heading_substring():
    path = _create_heading_docx()
    try:
        # Markdown mode
        result = get_content(path, heading="Section 1.1")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        assert meta["count"] == 1
        assert "Section 1.1 content" in meta["markdown"]

        # JSON mode
        result_json = get_content(path, heading="Section 1.1", format="json")
        meta_json, _ = _parse_content_result(result_json)
        assert meta_json["sections"][0]["title"] == "Section 1.1"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_by_heading_multiple():
    path = _create_heading_docx()
    try:
        result = get_content(path, heading="Section")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "error"
        assert "Multiple sections matched" in meta["message"]
        assert len(meta["candidates"]) >= 2
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_by_heading_not_found():
    path = _create_heading_docx()
    try:
        result = get_content(path, heading="NonExistent")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "error"
        assert "No section found" in meta["message"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_content_markdown_with_table():
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("Hardware Spec", level=1)
    doc.add_paragraph("Device pinout table:")
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Pin"
    t.rows[0].cells[1].text = "Name"
    t.rows[0].cells[2].text = "Type"
    t.rows[1].cells[0].text = "1"
    t.rows[1].cells[1].text = "VCC"
    t.rows[1].cells[2].text = "Power"
    doc.save(path)
    try:
        result = get_content(path, heading="Hardware Spec")
        meta, _ = _parse_content_result(result)
        assert meta["status"] == "ok"
        md = meta["markdown"]
        assert "# Hardware Spec" in md
        assert "Device pinout table:" in md
        assert "| Pin | Name | Type |" in md
        assert "| 1 | VCC | Power |" in md
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


# ── grep tests ──


def test_grep_basic():
    from carrot_mcp_office.word import grep
    path = _docx()
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Goodbye World")
        doc.add_paragraph("Hello Again")
        doc.save(path)

        result = grep(path, "Hello")
        assert result["status"] == "ok"
        assert result["count"] == 2
        assert result["matches"][0]["type"] == "paragraph"
        assert result["matches"][0]["index"] == 0
        assert result["matches"][0]["text"] == "Hello World"
        assert result["matches"][1]["index"] == 2
    finally:
        _cleanup(path)


def test_grep_case_insensitive():
    from carrot_mcp_office.word import grep
    path = _docx()
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.save(path)

        result = grep(path, "hello")
        assert result["status"] == "ok"
        assert result["count"] == 1
    finally:
        _cleanup(path)


def test_grep_regex():
    from carrot_mcp_office.word import grep
    path = _docx()
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("abc 123")
        doc.add_paragraph("xyz 456")
        doc.add_paragraph("abc 789")
        doc.save(path)

        result = grep(path, r"abc \d+", regex=True)
        assert result["status"] == "ok"
        assert result["count"] == 2
    finally:
        _cleanup(path)


def test_grep_regex_invalid():
    from carrot_mcp_office.word import grep
    path = _docx()
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("test")
        doc.save(path)

        result = grep(path, r"[invalid", regex=True)
        assert result["status"] == "error"
        assert "Invalid regex" in result["message"]
    finally:
        _cleanup(path)


def test_grep_context():
    from carrot_mcp_office.word import grep
    path = _docx()
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("line 0")
        doc.add_paragraph("line 1")
        doc.add_paragraph("TARGET")
        doc.add_paragraph("line 3")
        doc.add_paragraph("line 4")
        doc.save(path)

        result = grep(path, "TARGET")
        assert result["status"] == "ok"
        m = result["matches"][0]
        assert m["context_before"] == ["line 1"]
        assert m["context_after"] == ["line 3"]
    finally:
        _cleanup(path)


def test_grep_no_match():
    from carrot_mcp_office.word import grep
    path = _docx()
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("hello")
        doc.save(path)

        result = grep(path, "xyz")
        assert result["status"] == "ok"
        assert result["count"] == 0
        assert result["matches"] == []
    finally:
        _cleanup(path)


def test_grep_table_cell_penetration():
    """Test grep searching inside Word table cells."""
    from carrot_mcp_office.word import grep
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("Register Map", level=1)
    doc.add_heading("Control Registers", level=2)
    t = doc.add_table(rows=3, cols=4)
    # Header
    t.rows[0].cells[0].text = "Address"
    t.rows[0].cells[1].text = "Name"
    t.rows[0].cells[2].text = "R/W"
    t.rows[0].cells[3].text = "Description"
    # Row 1
    t.rows[1].cells[0].text = "0x00"
    t.rows[1].cells[1].text = "CTRL_REG"
    t.rows[1].cells[2].text = "RW"
    t.rows[1].cells[3].text = "System control register"
    # Row 2
    t.rows[2].cells[0].text = "0x04"
    t.rows[2].cells[1].text = "STATUS_REG"
    t.rows[2].cells[2].text = "RO"
    t.rows[2].cells[3].text = "Status flags"
    doc.save(path)
    try:
        result = grep(path, "CTRL_REG")
        assert result["status"] == "ok"
        assert result["count"] == 1
        m = result["matches"][0]
        assert m["type"] == "table"
        assert m["table_index"] == 0
        assert m["row_index"] == 1
        assert m["matched_col"] == 1
        assert m["matched_text"] == "CTRL_REG"
        assert m["header"] == ["Address", "Name", "R/W", "Description"]
        assert m["row_data"] == ["0x00", "CTRL_REG", "RW", "System control register"]
        assert m["heading_path"] == ["Register Map", "Control Registers"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_grep_mixed_paragraph_and_table():
    """Test grep finding matches in both paragraphs and tables."""
    from carrot_mcp_office.word import grep
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    doc.add_heading("Chapter 1: SPI Master", level=1)
    doc.add_paragraph("The SPI peripheral supports master mode with OP_MODE configuration.")
    t = doc.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Bit"
    t.rows[0].cells[1].text = "Symbol"
    t.rows[0].cells[2].text = "Description"
    t.rows[1].cells[0].text = "7:4"
    t.rows[1].cells[1].text = "OP_MODE"
    t.rows[1].cells[2].text = "Select operating mode"
    doc.save(path)
    try:
        result = grep(path, "OP_MODE")
        assert result["status"] == "ok"
        assert result["count"] == 2
        types = [m["type"] for m in result["matches"]]
        assert types == ["paragraph", "table"]
        # Both share the heading path
        assert result["matches"][0]["heading_path"] == ["Chapter 1: SPI Master"]
        assert result["matches"][1]["heading_path"] == ["Chapter 1: SPI Master"]
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_heading_level_chinese_and_english():
    from carrot_mcp_office.word import _heading_level
    assert _heading_level("Heading 1") == 1
    assert _heading_level("heading 2") == 2
    assert _heading_level("Heading3") == 3
    assert _heading_level("标题 1") == 1
    assert _heading_level("标题 2") == 2
    assert _heading_level("标题3") == 3
    assert _heading_level("标题 ４") == 4
    assert _heading_level("Normal") is None
    assert _heading_level("Title") is None
    assert _heading_level("") is None


def test_get_outline_chinese_styles():
    """Test get_outline parsing Chinese Word styles like '标题 1'."""
    from docx import Document
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    p1 = doc.add_paragraph("第一章 概述")
    p1.style = doc.styles["Heading 1"]
    p1.style.name = "标题 1"
    p2 = doc.add_paragraph("1.1 系统架构")
    p2.style = doc.styles["Heading 2"]
    p2.style.name = "标题 2"
    doc.save(path)
    try:
        result = get_outline(path)
        assert result["status"] == "ok"
        assert result["count"] == 2
        assert result["flat"][0]["title"] == "第一章 概述"
        assert result["flat"][0]["level"] == 1
        assert result["flat"][1]["title"] == "1.1 系统架构"
        assert result["flat"][1]["level"] == 2
        assert result["flat"][1]["parent"] == "第一章 概述"
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_get_outline_direct_outline_lvl_xml():
    """Test get_outline recognizing direct <w:outlineLvl> on paragraph properties."""
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    p = doc.add_paragraph("Custom Outline Level Paragraph")
    # Add direct <w:outlineLvl w:val="0"/> (Level 1)
    pPr = p._element.get_or_add_pPr()
    lvl_xml = parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="0"/>')
    pPr.append(lvl_xml)
    doc.save(path)
    try:
        result = get_outline(path)
        assert result["status"] == "ok"
        assert result["count"] == 1
        assert result["flat"][0]["title"] == "Custom Outline Level Paragraph"
        assert result["flat"][0]["level"] == 1
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_insert_para_multiline_and_list():
    """Test insert_para with multi-line string and list of strings."""
    from docx import Document
    path = _docx()
    try:
        # Multi-line string
        res = insert_para(path, "Line 1\nLine 2\nLine 3")
        assert res["status"] == "ok"
        assert res["inserted_count"] == 3
        doc = Document(path)
        assert len(doc.paragraphs) == 3
        assert [p.text for p in doc.paragraphs] == ["Line 1", "Line 2", "Line 3"]

        # List of strings
        res_list = insert_para(path, ["Line 4", "Line 5"])
        assert res_list["status"] == "ok"
        assert res_list["inserted_count"] == 2
        doc = Document(path)
        assert len(doc.paragraphs) == 5
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_format_para_without_initial_runs():
    """Test format_para applying styles even when paragraph has no runs initially."""
    from docx import Document
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    d = tempfile.mkdtemp(prefix="test_office_")
    path = os.path.join(d, "test.docx")
    doc = Document()
    p = doc.add_paragraph()
    r_elem = parse_xml(f'<w:r {nsdecls("w")}><w:t>Plain Text Without Runs</w:t></w:r>')
    p._element.append(r_elem)
    doc.save(path)
    try:
        res = format_para(path, index=0, bold=True, font_size=14, font_color="FF0000")
        assert res["status"] == "ok"
        doc = Document(path)
        p = doc.paragraphs[0]
        assert len(p.runs) >= 1
        assert p.runs[0].bold is True
        assert p.runs[0].font.color.rgb is not None
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)


def test_inspect_slimming():
    """Test inspect returns clean metadata without paragraphs list."""
    path = _docx()
    try:
        insert_para(path, "Paragraph 1\nParagraph 2")
        insert_table(path, 2, 3)
        res = word_inspect(path)
        assert res["status"] == "ok"
        assert res["total_paragraphs"] == 2
        assert res["non_empty_paragraphs"] == 2
        assert res["table_count"] == 1
        assert res["tables"][0] == {"index": 0, "rows": 2, "cols": 3}
        assert "paragraphs" not in res
    finally:
        _cleanup(path)
        if os.path.exists(path):
            os.unlink(path)

