"""Word tools for carrot-mcp-office using python-docx."""

from __future__ import annotations

import base64
import json
import os
import re

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mcp.types import ImageContent, TextContent

from carrot_mcp_office._mcp import mcp, _save_and_return
from carrot_mcp_office._format import render_table_markdown, parse_index_range
from carrot_mcp_office.convert import ensure_docx_format

_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
_A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"

_HEADING_PATTERN = re.compile(r"^(?:heading|标题)\s*([1-9]|１|２|３|４|５|６|７|８|９)$", re.IGNORECASE)
_FULLWIDTH_DIGIT_MAP = {
    "１": 1, "２": 2, "３": 3, "４": 4, "５": 5,
    "６": 6, "７": 7, "８": 8, "９": 9,
}


def _open_or_create_document(path: str) -> Document:
    """Open existing document or create a new one."""
    if os.path.exists(path):
        return Document(path)
    doc = Document()
    doc.save(path)
    return Document(path)


def _handle_docx(path: str) -> tuple[str, dict | None]:
    """Ensure docx format, return (resolved_path, error_or_none)."""
    resolved, err = ensure_docx_format(path)
    if err:
        return path, {"status": "error", "message": err}
    return resolved, None


@mcp.tool()
def inspect(path: str) -> dict:
    """Inspect document structure (paragraphs, tables, images, styles).

    Args:
        path: Absolute path to the .doc/.docx file.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        non_empty_paras = sum(1 for p in doc.paragraphs if p.text.strip())
        tables = [{"index": i, "rows": len(t.rows), "cols": len(t.columns)} for i, t in enumerate(doc.tables)]
        image_count = 0
        for p in doc.paragraphs:
            for run in p.runs:
                if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
                    image_count += 1
        styles = list(dict.fromkeys(p.style.name for p in doc.paragraphs))
        return {
            "status": "ok",
            "path": path,
            "total_paragraphs": len(doc.paragraphs),
            "non_empty_paragraphs": non_empty_paras,
            "table_count": len(doc.tables),
            "image_count": image_count,
            "styles_used": styles,
            "tables": tables,
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def insert_para(path: str, text: str | list[str], index: int | None = None) -> dict:
    """Insert one or more paragraphs at the specified position.

    If text contains newlines ('\\n') or is passed as a list of strings,
    it will be inserted as multiple consecutive paragraphs in visual order.

    Args:
        path: Absolute path to the .doc/.docx file.
        text: Paragraph text (multi-line string or list of strings).
        index: Position (0-based) to insert at. None appends at end.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = _open_or_create_document(path)

        if isinstance(text, list):
            lines = [line for item in text for line in str(item).split("\n")]
        else:
            lines = str(text).split("\n")

        if not lines:
            lines = [""]

        inserted_indices = []
        if index is None or index >= len(doc.paragraphs):
            for line in lines:
                doc.add_paragraph(line)
                inserted_indices.append(len(doc.paragraphs) - 1)
        else:
            ref_para = doc.paragraphs[index]
            for line_offset, line in enumerate(lines):
                new_para = doc.add_paragraph(line)
                ref_para._element.addprevious(new_para._element)
                inserted_indices.append(index + line_offset)

        doc.save(path)
        return _save_and_return(
            path,
            "insert_para",
            {
                "status": "ok",
                "text": text,
                "inserted_count": len(lines),
                "index": inserted_indices[0] if inserted_indices else 0,
                "indices": inserted_indices,
            },
        )
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def modify_para(path: str, index: int, text: str) -> dict:
    """Modify an existing paragraph's text.

    Args:
        path: Absolute path to the .doc/.docx file.
        index: Paragraph index (0-based).
        text: New text content.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"status": "error", "message": f"Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})"}
        para = doc.paragraphs[index]
        old_text = para.text
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = text
        else:
            para.text = text
        doc.save(path)
        return _save_and_return(path, "modify_para", {"status": "ok", "index": index, "old_text": old_text, "new_text": text})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def format_para(
    path: str,
    index: int,
    style: str | None = None,
    alignment: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    font_size: int | None = None,
    font_color: str | None = None,
) -> dict:
    """Format a paragraph.

    Args:
        path: Absolute path to the .doc/.docx file.
        index: Paragraph index (0-based).
        style: Style name (e.g. "Heading 1", "Normal", "Title").
        alignment: Text alignment (left, center, right, justify).
        bold: Bold text.
        italic: Italic text.
        font_size: Font size in points.
        font_color: Font color as hex string (e.g. "FF0000").
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"status": "error", "message": f"Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})"}
        para = doc.paragraphs[index]
        if style:
            try:
                para.style = doc.styles[style]
            except KeyError:
                return {"status": "error", "message": f"Style '{style}' not found"}
        if alignment:
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if alignment not in align_map:
                return {"status": "error", "message": f"Invalid alignment '{alignment}'. Use: left, center, right, justify"}
            para.alignment = align_map[alignment]
        # Ensure paragraph has at least one run to receive formatting
        if not para.runs and para.text:
            raw_text = para.text
            para.text = ""
            para.add_run(raw_text)
        elif not para.runs and (bold is not None or italic is not None or font_size or font_color):
            para.add_run("")

        for run in para.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if font_size:
                run.font.size = Inches(font_size / 72)
            if font_color:
                run.font.color.rgb = RGBColor.from_string(font_color)
        doc.save(path)
        return _save_and_return(path, "format_para", {"status": "ok", "index": index})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def delete_para(path: str, index: int) -> dict:
    """Delete a paragraph.

    Args:
        path: Absolute path to the .doc/.docx file.
        index: Paragraph index (0-based).
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"status": "error", "message": f"Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})"}
        para = doc.paragraphs[index]
        parent = para._element.getparent()
        parent.remove(para._element)
        doc.save(path)
        return _save_and_return(path, "delete_para", {"status": "ok", "index": index})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def insert_table(path: str, rows: int, cols: int, data: list[list] | None = None, index: int | None = None) -> dict:
    """Insert a table.

    Args:
        path: Absolute path to the .doc/.docx file.
        rows: Number of rows.
        cols: Number of columns.
        data: Optional 2D array of cell values.
        index: Position (0-based) to insert at. None appends at end.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = _open_or_create_document(path)
        if data and (len(data) != rows or any(len(row) != cols for row in data)):
            return {"status": "error", "message": "Data dimensions don't match rows/cols"}
        ref_table = None
        if index is not None and index < len(doc.tables):
            ref_table = doc.tables[index]
        table = doc.add_table(rows=rows, cols=cols)
        if data:
            for r_idx, row in enumerate(data):
                for c_idx, val in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = str(val) if val is not None else ""
        if ref_table is not None:
            ref_table._element.addprevious(table._element)
        doc.save(path)
        return _save_and_return(path, "insert_table", {"status": "ok", "rows": rows, "cols": cols, "index": index if index is not None else len(doc.tables) - 1})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def modify_table(path: str, table_index: int, row: int, col: int, text: str) -> dict:
    """Modify a table cell.

    Args:
        path: Absolute path to the .doc/.docx file.
        table_index: Table index (0-based).
        row: Row index (0-based).
        col: Column index (0-based).
        text: New cell text.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if table_index < 0 or table_index >= len(doc.tables):
            return {"status": "error", "message": f"Table index {table_index} out of range (0-{len(doc.tables)-1})"}
        table = doc.tables[table_index]
        if row < 0 or row >= len(table.rows):
            return {"status": "error", "message": f"Row {row} out of range (0-{len(table.rows)-1})"}
        if col < 0 or col >= len(table.columns):
            return {"status": "error", "message": f"Column {col} out of range (0-{len(table.columns)-1})"}
        old_text = table.rows[row].cells[col].text
        table.rows[row].cells[col].text = text
        doc.save(path)
        return _save_and_return(path, "modify_table", {"status": "ok", "table_index": table_index, "row": row, "col": col, "old_text": old_text, "new_text": text})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def format_table(path: str, table_index: int, style: str | None = None) -> dict:
    """Format a table.

    Args:
        path: Absolute path to the .doc/.docx file.
        table_index: Table index (0-based).
        style: Table style name (e.g. "Table Grid", "Light Shading").
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if table_index < 0 or table_index >= len(doc.tables):
            return {"status": "error", "message": f"Table index {table_index} out of range (0-{len(doc.tables)-1})"}
        table = doc.tables[table_index]
        if style:
            try:
                table.style = doc.styles[style]
            except KeyError:
                return {"status": "error", "message": f"Style '{style}' not found"}
        doc.save(path)
        return _save_and_return(path, "format_table", {"status": "ok", "table_index": table_index, "style": style})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def delete_table(path: str, table_index: int) -> dict:
    """Delete a table.

    Args:
        path: Absolute path to the .doc/.docx file.
        table_index: Table index (0-based).
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if table_index < 0 or table_index >= len(doc.tables):
            return {"status": "error", "message": f"Table index {table_index} out of range (0-{len(doc.tables)-1})"}
        table = doc.tables[table_index]
        parent = table._element.getparent()
        parent.remove(table._element)
        doc.save(path)
        return _save_and_return(path, "delete_table", {"status": "ok", "table_index": table_index})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def get_table(
    path: str,
    table_index: int,
    row_range: list | int | str | None = None,
    format: str = "markdown",
) -> dict:
    """Read table content as Markdown table or 2D array, with optional row filtering.

    Args:
        path: Absolute path to the .doc/.docx file.
        table_index: Table index (0-based).
        row_range: Optional row indices or range to slice (e.g. "0-10", "0, 5-8", [0, 2]).
                   If row 0 (header) is present and not explicitly included, it is
                   automatically included at the top for context in markdown mode.
        format: Output format, either "markdown" (default) or "json".
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if table_index < 0 or table_index >= len(doc.tables):
            return {"status": "error", "message": f"Table index {table_index} out of range (0-{len(doc.tables)-1})"}
        table = doc.tables[table_index]
        all_data = [[cell.text for cell in row.cells] for row in table.rows]

        if row_range is not None:
            selected_rows = parse_index_range(row_range, len(all_data) - 1)
            if 0 not in selected_rows and len(all_data) > 0 and format == "markdown":
                selected_rows = [0] + selected_rows
            data = [all_data[i] for i in selected_rows if 0 <= i < len(all_data)]
        else:
            data = all_data

        res = {
            "status": "ok",
            "path": path,
            "table_index": table_index,
            "total_rows": len(all_data),
            "rows": len(data),
            "cols": len(data[0]) if data else 0,
        }
        if format == "markdown":
            res["markdown"] = render_table_markdown(data)
        else:
            res["data"] = data
        return res
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def insert_image(path: str, image_path: str, index: int | None = None, width: float | None = None) -> dict:
    """Insert an image into the document.

    Args:
        path: Absolute path to the .doc/.docx file.
        image_path: Absolute path to the image file.
        index: Paragraph position (0-based) to insert at. None appends at end.
        width: Image width in inches. None uses original size.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = _open_or_create_document(path)
        if index is None or index >= len(doc.paragraphs):
            para = doc.add_paragraph()
        else:
            ref_para = doc.paragraphs[index]
            para = doc.add_paragraph()
            ref_para._element.addprevious(para._element)
        run = para.add_run()
        if width:
            run.add_picture(image_path, width=Inches(width))
        else:
            run.add_picture(image_path)
        doc.save(path)
        return _save_and_return(path, "insert_image", {"status": "ok", "image_path": image_path, "index": index})
    except Exception as e:
        return {"status": "error", "message": str(e)}


@mcp.tool()
def delete_image(path: str, image_index: int) -> dict:
    """Delete an inline image by its occurrence index.

    Args:
        path: Absolute path to the .doc/.docx file.
        image_index: Image occurrence index (0-based, counting across all paragraphs).
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        current_index = 0
        for para in doc.paragraphs:
            for run in para.runs:
                drawings = run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")
                for drawing in drawings:
                    if current_index == image_index:
                        parent = drawing.getparent()
                        parent.remove(drawing)
                        doc.save(path)
                        return _save_and_return(path, "delete_image", {"status": "ok", "image_index": image_index})
                    current_index += 1
        return {"status": "error", "message": f"Image index {image_index} out of range (found {current_index} images)"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def _heading_level(style_name: str) -> int | None:
    """Return heading level (1-9) from style name, or None if not a heading.

    Supports English ('Heading 1', 'heading 1', 'Heading1') and
    Chinese ('标题 1', '标题1', '标题 １') style names.
    """
    if not style_name:
        return None
    m = _HEADING_PATTERN.match(style_name.strip())
    if m:
        digit_str = m.group(1)
        if digit_str in _FULLWIDTH_DIGIT_MAP:
            return _FULLWIDTH_DIGIT_MAP[digit_str]
        try:
            return int(digit_str)
        except ValueError:
            return None
    return None


def _build_style_outline_map(doc: Document) -> dict[str, int]:
    """Map style names and style IDs to their effective outline level (1-9).

    Scans doc.styles for explicit <w:outlineLvl> elements in style definitions,
    and falls back to standard localized heading patterns.
    """
    style_map: dict[str, int] = {}
    try:
        for s in doc.styles:
            # 1. Check if style XML has explicit <w:outlineLvl>
            try:
                pPr = s._element.pPr
                if pPr is not None:
                    lvl_elem = pPr.find(f".//{_NS}outlineLvl")
                    if lvl_elem is not None:
                        val = lvl_elem.get(f"{_NS}val")
                        if val is not None and val.isdigit():
                            lvl = int(val) + 1
                            if 1 <= lvl <= 9:
                                style_map[s.name] = lvl
                                if s.style_id:
                                    style_map[s.style_id] = lvl
                                continue
            except Exception:
                pass

            # 2. Check localized / built-in heading name pattern
            lvl = _heading_level(s.name) or (_heading_level(s.style_id) if s.style_id else None)
            if lvl is not None:
                style_map[s.name] = lvl
                if s.style_id:
                    style_map[s.style_id] = lvl
    except Exception:
        pass
    return style_map


def _get_paragraph_outline_level(para, style_map: dict[str, int]) -> int | None:
    """Get the effective outline level (1-9) for a paragraph.

    1. Checks direct paragraph property override (<w:pPr><w:outlineLvl w:val="N"/></w:pPr>).
    2. Checks inherited style outline level from styles.xml (via style_map).
    3. Falls back to style name pattern matching.
    """
    # 1. Direct paragraph property override
    try:
        pPr = para._element.pPr
        if pPr is not None:
            lvl_elem = pPr.find(f".//{_NS}outlineLvl")
            if lvl_elem is not None:
                val = lvl_elem.get(f"{_NS}val")
                if val is not None and val.isdigit():
                    lvl = int(val) + 1
                    if 1 <= lvl <= 9:
                        return lvl
    except Exception:
        pass

    # 2. Inherited from style map
    if para.style:
        if para.style.name in style_map:
            return style_map[para.style.name]
        if para.style.style_id in style_map:
            return style_map[para.style.style_id]
        lvl = _heading_level(para.style.name) or (_heading_level(para.style.style_id) if para.style.style_id else None)
        if lvl is not None:
            return lvl

    return None


def _flatten_outline(nodes: list[dict]) -> list[dict]:
    """Flatten outline tree into a list with parent tracking."""
    result = []
    for node in nodes:
        result.append({
            "level": node["level"],
            "title": node["title"],
            "index": node["index"],
            "parent": node.get("parent"),
        })
        if node.get("children"):
            result.extend(_flatten_outline(node["children"]))
    return result


def _has_images_in_para(para) -> bool:
    return bool(para._element.findall(f".//{_NS}drawing"))


def _extract_images_from_para(para) -> list[tuple[bytes, str]]:
    """Extract (image_bytes, mime_type) from a paragraph's inline drawings."""
    images = []
    for drawing in para._element.findall(f".//{_NS}drawing"):
        for blip in drawing.findall(f".//{_A_NS}blip"):
            rId = blip.get(f"{_R_NS}embed")
            if not rId:
                continue
            rel = para.part.rels.get(rId)
            if rel is None:
                continue
            img_part = rel.target_part
            img_bytes = img_part.blob
            content_type = img_part.content_type or "image/png"
            images.append((img_bytes, content_type))
    return images


_parse_sections = parse_index_range


@mcp.tool()
def get_outline(path: str) -> dict:
    """Get document outline (heading hierarchy).

    Returns both a hierarchical tree and a flat list of headings (Heading 1-9).
    Each node contains: level, title, index (paragraph position), parent.

    To fetch section content, pass the flat array's 0-based position indices
    to get_content. For example, if flat returns:
      [{"level":1,"title":"Intro","index":0}, {"level":2,"title":"A","index":3}]
    Then sections=[0, 1] fetches "Intro" and "A".

    Args:
        path: Absolute path to the .doc/.docx file.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        style_map = _build_style_outline_map(doc)
        headings = []
        for i, para in enumerate(doc.paragraphs):
            level = _get_paragraph_outline_level(para, style_map)
            if level is not None:
                headings.append({"level": level, "title": para.text, "index": i})

        stack: list[dict] = []
        tree: list[dict] = []
        for h in headings:
            node = {"level": h["level"], "title": h["title"], "index": h["index"], "children": []}
            while stack and stack[-1]["level"] >= h["level"]:
                stack.pop()
            if stack:
                node["parent"] = stack[-1]["title"]
                stack[-1]["children"].append(node)
            else:
                tree.append(node)
            stack.append(node)

        flat = _flatten_outline(tree)
        return {
            "status": "ok",
            "path": path,
            "outline": tree,
            "flat": flat,
            "count": len(flat),
        }
    except Exception as e:
        return {"status": "error", "message": str(e)}


def _get_element_heading_paths(doc: Document) -> tuple[dict[int, list[str]], dict[int, list[str]]]:
    """Map paragraph indices and table indices to their hierarchical heading path.

    Returns:
        (para_heading_paths, table_heading_paths)
        where para_heading_paths is {para_index: ["Heading 1", "Heading 2", ...]}
        and table_heading_paths is {table_index: ["Heading 1", "Heading 2", ...]}
    """
    p_map = {p._element: i for i, p in enumerate(doc.paragraphs)}
    t_map = {t._element: i for i, t in enumerate(doc.tables)}
    style_map = _build_style_outline_map(doc)

    para_paths: dict[int, list[str]] = {}
    table_paths: dict[int, list[str]] = {}

    heading_stack: list[dict] = []  # [{"level": int, "title": str}]

    for child in doc.element.body:
        if child in p_map:
            p_idx = p_map[child]
            p = doc.paragraphs[p_idx]
            level = _get_paragraph_outline_level(p, style_map)
            if level is not None and p.text.strip():
                while heading_stack and heading_stack[-1]["level"] >= level:
                    heading_stack.pop()
                heading_stack.append({"level": level, "title": p.text.strip()})
            para_paths[p_idx] = [h["title"] for h in heading_stack]
        elif child in t_map:
            t_idx = t_map[child]
            table_paths[t_idx] = [h["title"] for h in heading_stack]

    return para_paths, table_paths


@mcp.tool()
def get_content(
    path: str,
    section: list | int | str | None = None,
    paragraph: list | int | str | None = None,
    heading: str | None = None,
    format: str = "markdown",
) -> list:
    """Get content for specific outline sections, paragraphs, or heading titles.

    Use get_outline or heading names to fetch section content including
    paragraphs, tables, and images.
    Images are returned as ImageContent attachments (not embedded in JSON).

    Args:
        path: Absolute path to the .doc/.docx file.
        section: Indices into the `flat` array returned by get_outline. Accepts:
            - int: single index, e.g. 0
            - str: range or comma-separated list, e.g. "0-9" or "0-4,6,8"
            - list: array of int/str, e.g. [0, 2, 5] or ["0-4,6,8"]
        paragraph: Direct paragraph indices (0-based position in the document).
            Accepts same format as section.
        heading: Heading title to search and fetch content for. Performs exact
            match first, then case-insensitive substring match.
        format: Output presentation format:
            - "markdown" (default): Clean high-density Markdown text for reading.
            - "json": Full structured metadata (indices, styles, ranges) for editing.

    Returns:
        list[TextContent | ImageContent] — first element is JSON metadata
        containing Markdown content or structured section data, followed by image attachments.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return [TextContent(type="text", text=json.dumps(err))]
        mode_count = sum(1 for m in (section, paragraph, heading) if m is not None)
        if mode_count != 1:
            return [TextContent(type="text", text=json.dumps({"status": "error", "message": "Must provide exactly one of 'section', 'paragraph', or 'heading'"}))]

        doc = Document(path)
        result: list = []

        if paragraph is not None:
            para_indices = _parse_sections(paragraph, len(doc.paragraphs) - 1)
            paragraphs = []
            para_texts_for_md = []
            img_idx = 0
            out_of_range = []
            for p_idx in para_indices:
                if p_idx < 0 or p_idx >= len(doc.paragraphs):
                    out_of_range.append(p_idx)
                    continue
                para = doc.paragraphs[p_idx]
                if para.text.strip():
                    para_texts_for_md.append(para.text.strip())
                    paragraphs.append({"index": p_idx, "text": para.text, "style": para.style.name})
                for img_bytes, mime in _extract_images_from_para(para):
                    result.append(ImageContent(
                        type="image",
                        data=base64.b64encode(img_bytes).decode(),
                        mimeType=mime,
                        context=f"Paragraph {p_idx}, image {img_idx}",
                    ))
                    img_idx += 1
            meta: dict = {
                "status": "ok",
                "path": path,
                "mode": "paragraph",
                "count": len(paragraphs),
            }
            if format == "json":
                meta["paragraphs"] = paragraphs
            else:
                meta["markdown"] = "\n\n".join(para_texts_for_md)
            if out_of_range:
                meta["warning"] = f"Paragraphs out of range (0-{len(doc.paragraphs)-1}): {out_of_range}"
            result.insert(0, TextContent(type="text", text=json.dumps(meta, ensure_ascii=False)))
            return result

        style_map = _build_style_outline_map(doc)
        headings = []
        for i, para in enumerate(doc.paragraphs):
            level = _get_paragraph_outline_level(para, style_map)
            if level is not None:
                headings.append({"level": level, "title": para.text, "index": i})

        stack: list[dict] = []
        tree: list[dict] = []
        for h in headings:
            node = {"level": h["level"], "title": h["title"], "index": h["index"], "children": []}
            while stack and stack[-1]["level"] >= h["level"]:
                stack.pop()
            if stack:
                node["parent"] = stack[-1]["title"]
                stack[-1]["children"].append(node)
            else:
                tree.append(node)
            stack.append(node)

        flat = _flatten_outline(tree)

        if heading is not None:
            query = heading.strip().lower()
            exact_matches = [i for i, node in enumerate(flat) if node["title"].strip().lower() == query]
            if len(exact_matches) == 1:
                sec_indices = [exact_matches[0]]
            elif len(exact_matches) > 1:
                sec_indices = exact_matches
            else:
                sub_matches = [i for i, node in enumerate(flat) if query in node["title"].strip().lower()]
                if len(sub_matches) == 1:
                    sec_indices = [sub_matches[0]]
                elif len(sub_matches) > 1:
                    candidates = [{"section": i, "title": flat[i]["title"], "level": flat[i]["level"]} for i in sub_matches]
                    return [TextContent(type="text", text=json.dumps({
                        "status": "error",
                        "message": f"Multiple sections matched heading '{heading}'. Please specify exact title or section index.",
                        "candidates": candidates,
                    }, ensure_ascii=False))]
                else:
                    return [TextContent(type="text", text=json.dumps({
                        "status": "error",
                        "message": f"No section found matching heading '{heading}'",
                    }, ensure_ascii=False))]
        else:
            sec_indices = _parse_sections(section, len(flat) - 1)

        sections_meta = []
        section_markdowns = []
        out_of_range = []

        p_elem_map = {p._element: (i, p) for i, p in enumerate(doc.paragraphs)}
        t_elem_map = {t._element: (i, t) for i, t in enumerate(doc.tables)}

        for sec_idx in sec_indices:
            if sec_idx < 0 or sec_idx >= len(flat):
                out_of_range.append(sec_idx)
                continue

            sec_node = flat[sec_idx]
            start_p_idx = sec_node["index"]
            start_elem = doc.paragraphs[start_p_idx]._element

            end_p_idx = len(doc.paragraphs)
            end_elem = None
            for k in range(sec_idx + 1, len(flat)):
                if flat[k]["level"] <= sec_node["level"]:
                    end_p_idx = flat[k]["index"]
                    end_elem = doc.paragraphs[end_p_idx]._element
                    break

            paragraphs = []
            tables = []
            sec_md_parts = []
            prefix = "#" * min(6, sec_node["level"])
            sec_md_parts.append(f"{prefix} {sec_node['title']}")
            img_idx = 0

            # 1. Collect structured paragraphs
            for j in range(start_p_idx, end_p_idx):
                para = doc.paragraphs[j]
                if para.text.strip():
                    paragraphs.append({"index": j, "text": para.text, "style": para.style.name})
                for img_bytes, mime in _extract_images_from_para(para):
                    result.append(ImageContent(
                        type="image",
                        data=base64.b64encode(img_bytes).decode(),
                        mimeType=mime,
                        context=f"Section {sec_idx} ({sec_node['title']}), image {img_idx}",
                    ))
                    img_idx += 1

            # 2. Collect body elements in visual order for markdown and tables
            in_section = False
            for child in doc.element.body:
                if child is start_elem:
                    in_section = True
                    continue
                if end_elem is not None and child is end_elem:
                    break
                if not in_section:
                    continue

                if child in p_elem_map:
                    _, para = p_elem_map[child]
                    if para.text.strip():
                        sec_md_parts.append(para.text.strip())
                elif child in t_elem_map:
                    t_idx, tbl = t_elem_map[child]
                    rows = [[cell.text for cell in row.cells] for row in tbl.rows]
                    tables.append({"table_index": t_idx, "rows": len(rows), "cols": len(rows[0]) if rows else 0, "data": rows})
                    tbl_md = render_table_markdown(rows)
                    if tbl_md:
                        sec_md_parts.append(tbl_md)

            sec_entry: dict = {
                "section": sec_idx,
                "title": sec_node["title"],
                "level": sec_node["level"],
                "paragraph_range": [start_p_idx, end_p_idx - 1],
                "image_count": img_idx,
                "paragraphs": paragraphs,
                "tables": tables,
            }
            sections_meta.append(sec_entry)
            section_markdowns.append("\n\n".join(sec_md_parts))

        meta = {
            "status": "ok",
            "path": path,
            "mode": "heading" if heading is not None else "section",
            "count": len(sec_indices),
        }
        if format == "json":
            meta["sections"] = sections_meta
        else:
            meta["markdown"] = "\n\n---\n\n".join(section_markdowns)

        if out_of_range:
            meta["warning"] = f"Sections out of range (0-{len(flat)-1}): {out_of_range}"
        result.insert(0, TextContent(type="text", text=json.dumps(meta, ensure_ascii=False)))
        return result
    except Exception as e:
        return [TextContent(type="text", text=json.dumps({"status": "error", "message": str(e)}))]


@mcp.tool()
def grep(path: str, pattern: str, regex: bool = False) -> dict:
    """Search for exact substring or regex in paragraphs and tables.

    Searches across all document paragraphs and table cells.
    Returns matched items with their enclosing heading path (chapter hierarchy),
    table headers, and surrounding context.

    Args:
        path: Absolute path to the .doc/.docx file.
        pattern: Substring to match (case-insensitive) or regex pattern.
        regex: If true, treat pattern as a Python regular expression.
    """
    try:
        path, err = _handle_docx(path)
        if err:
            return err
        doc = Document(path)
        if regex:
            regex_pattern = re.compile(pattern, re.IGNORECASE)
            def _match(text: str) -> bool:
                return bool(regex_pattern.search(text))
        else:
            def _match(text: str) -> bool:
                return pattern.lower() in text.lower()

        para_paths, table_paths = _get_element_heading_paths(doc)
        matches = []

        # 1. Search paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() and _match(para.text):
                ctx_before = [doc.paragraphs[j].text for j in range(max(0, i - 1), i) if doc.paragraphs[j].text.strip()]
                ctx_after = [doc.paragraphs[j].text for j in range(i + 1, min(len(doc.paragraphs), i + 2)) if doc.paragraphs[j].text.strip()]
                matches.append({
                    "type": "paragraph",
                    "index": i,
                    "text": para.text,
                    "style": para.style.name,
                    "heading_path": para_paths.get(i, []),
                    "context_before": ctx_before,
                    "context_after": ctx_after,
                })

        # 2. Search tables
        for t_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue
            header = [cell.text.strip() for cell in table.rows[0].cells]
            t_heading_path = table_paths.get(t_idx, [])

            for r_idx, row in enumerate(table.rows):
                row_texts = [cell.text.strip() for cell in row.cells]
                matched_cols = [c_idx for c_idx, cell_text in enumerate(row_texts) if _match(cell_text)]
                if matched_cols:
                    matches.append({
                        "type": "table",
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "matched_col": matched_cols[0],
                        "matched_text": row_texts[matched_cols[0]],
                        "heading_path": t_heading_path,
                        "header": header,
                        "row_data": row_texts,
                    })

        return {
            "status": "ok",
            "path": path,
            "pattern": pattern,
            "regex": regex,
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "matches": matches,
            "count": len(matches),
        }
    except re.error as e:
        return {"status": "error", "message": f"Invalid regex: {e}"}
    except Exception as e:
        return {"status": "error", "message": str(e)}
